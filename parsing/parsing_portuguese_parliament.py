import pathlib
import datetime
from docx import Document
from io import StringIO
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser
from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer
import re
import pandas as pd
import numpy as np
import math
from eia_crawling.spiders.utils import get_parliament_name, get_iso_2_digit_code, get_iso_3_digit_code, write_csv, \
    normalize_string

COUNTRY = 'portugal'

CLOSING_PHRASE = 'encerrada a sessão|' \
                 'terminado este debate|' \
                 '(E|e)stá encerrada sessão|' \
                 'concluídos os nossos trabalhos|' \
                 'terminada a sessão|' \
                 'encerrada a reunião|' \
                 '(D|d)eclaro encerrada|' \
                 'terminámos nos nossos trabalhos|' \
                 'estão encerrados os nossos trabalhos|' \
                 'amanhã não há reunião plenária em virtude da realização do Congresso do Partido Socialista, a quem desejo o maior sucesso.'

MONTHS_DCT = {'Janeiro': 1,
              'Fevereiro': 2,
              'Março': 3,
              'Abril': 4,
              'Maio': 5,
              'Junho': 6,
              'Julho': 7,
              'Agosto': 8,
              'Setembro': 9,
              'Outubro': 10,
              'Novembro': 11,
              'Novemro': 11,
              'Dezembro': 12}

SPEAKER_ROLES = re.compile(r'(Secretário|Presidente|Secretária|Ministro)')

output_string = StringIO()


def parse_portuguese_parliament(source_pdf_path: pathlib,
                                year_path: pathlib.Path,
                                year: int):
    with open(str(source_pdf_path), 'rb') as in_file:
        parser = PDFParser(in_file)
        doc = PDFDocument(parser)
        rsrcmgr = PDFResourceManager()
        device = TextConverter(rsrcmgr, output_string, laparams=LAParams())
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        for page in PDFPage.create_pages(doc):
            interpreter.process_page(page)

    doc_lines = []

    for page_layout in extract_pages(str(source_pdf_path)):
        for element in page_layout:
            if isinstance(element, LTTextContainer):
                for line in element:
                    text_line = line.get_text()
                    text_line = text_line.strip('\n')
                    bbox = line.bbox
                    doc_lines.append([text_line, bbox])

    # Get speech breakers
    def get_speech_breakers(path: pathlib):
        """
        Converts .pdf document to .docx to extracts parts of text formatted in italics.
        :param path: path to .pdf file
        :returns dct
        """
        docx_filename = path.stem + '.docx'
        docx_path = path.parent.joinpath(docx_filename)

        # Extract text in italics
        breaking_lines = []
        doc = Document(docx_path)
        for p in doc.paragraphs:
            for run in p.runs:
                if run.italic:
                    breaking_lines.append(run.text.strip())

        # Filter sentences to get only those with at least 1 word
        breaking_lines = [sent for sent in breaking_lines if len(sent.split()) >= 1]
        # Leave 'Srs. Deputados presentes à sessão'
        breaking_lines = [sent for sent in breaking_lines if not re.match(r'Srs. Deputados presentes à sessão', sent)]
        # Remove punctuation from breaking lines
        breaking_lines = [re.sub(r'[^\w\s]', '', sent) for sent in breaking_lines]
        return breaking_lines

    # Filter out speech breakers
    try:
        breaking_lines = get_speech_breakers(source_pdf_path)
    except:
        breaking_lines = []
        print('DOCUMENT NOT PROCESSED')

    doc_lines = [i for i in doc_lines if re.sub(r'[^\w\s]', '', i[0].strip()) not in breaking_lines]

    # Filter out lines with page numbers
    page_pattern = re.compile(r'^\d+$')
    doc_lines = [i for i in doc_lines if not re.match(page_pattern, i[0].strip())]

    # Parse date,then filter out lines with date
    date_pattern = re.compile(r'^\d+\s+DE\s+'.format('|'.join(map(lambda x: x.upper(), MONTHS_DCT.keys()))))
    date_parsed = [i for i in doc_lines if re.match(date_pattern, i[0].strip())]

    try:
        date_parsed = date_parsed[0][0]
        date_parsed = date_parsed.replace(' DE ', ' ')
        date_day = int(date_parsed.split()[0])
        date_month = MONTHS_DCT[date_parsed.split()[1].capitalize()]
        date_year = int(date_parsed.split()[2])
        date = datetime.datetime(day=date_day, month=date_month, year=date_year)
        date = datetime.datetime.strftime(date, '%d.%m.%Y')
    except IndexError:
        print(f' Doc: {source_pdf_path.stem} Date not found')
        date = datetime.datetime(day=1, month=1, year=9999)
        date = datetime.datetime.strftime(date, '%d.%m.%Y')

    doc_lines = [i for i in doc_lines if not re.match(date_pattern, i[0].strip())]

    # Delete lines before the first speech
    speaker_pattern = re.compile(r'^(O(\s)+Sr.|A(\s)+Sr.([aª])|Sr.)\s+.*(\.|:)\s{0,3}—')
    for line in doc_lines:
        if re.match(speaker_pattern, line[0]):
            idx = doc_lines.index(line)
            doc_lines = doc_lines[idx:]
            break
    # Filter out empty lines:
    empty_pattern = re.compile(r'^\s*$')
    doc_lines = [i for i in doc_lines if not re.match(empty_pattern, i[0])]

    # Filter out title, e.g. 'I SÉRIE — NÚMERO 1'
    doc_lines = [i for i in doc_lines if 'SÉRIE —' not in i[0].strip()]

    # Filter out crowd reactions, e.g. 'Vozes do PS:', fix speech continuity
    crowd_pattern = re.compile(r'^Vozes\s{0,3}.*:')
    doc_lines = [i for i in doc_lines if not bool(re.match(crowd_pattern, i[0]))]

    # Delete information about participants in session. Delete text from 'Srs. Deputados presentes à sessão:' till next speaker
    try:
        start_of_list = next(
            i for i in doc_lines if re.match(r'Srs. Deputados presentes à sessão|Deputados presentes à sessão', i[0]))
        start_of_list_idx = doc_lines.index(start_of_list)
    except StopIteration:
        start_of_list = next(i for i in doc_lines if re.match(speaker_pattern, i[0]))
        start_of_list_idx = doc_lines.index(start_of_list)

    for i in range(start_of_list_idx + 1, len(doc_lines)):
        if re.match(r'^O(\s)+Sr.|A(\s)+Sr.ª\s+.*(:|\.)', doc_lines[i][0]):
            end_of_list_idx = i
            doc_lines = doc_lines[:start_of_list_idx] + doc_lines[end_of_list_idx:]
            break

    ## Find start lines of speeches
    i = 1
    for line in doc_lines:
        if re.match(speaker_pattern, line[0]):
            line.append(i)
            i += 1
        else:
            line.append(np.NaN)

    ## Find start lines of paragraphs
    # Find the standard location of indent
    indent = next(i for i in doc_lines if i[2] is not None)
    indent = indent[1][0]
    j = 1
    current_speech_number = 1
    for line in doc_lines:
        speech_number = line[2]
        if line[1][0] == indent:
            if speech_number > current_speech_number:
                j = 1
                line.append(j)
            else:
                line.append(j)
            j += 1
        else:
            line.append(np.NaN)

        if not math.isnan(speech_number):
            current_speech_number = speech_number

    # Get Speaker's name and party
    for line in doc_lines:
        if not math.isnan(line[2]):
            speaker_meta = re.split(r'(\.|:)\s+—', line[0], 1)[0].strip()
            speaker_meta = re.sub(r'O(\s)+Sr.|A(\s)+Sr.ª', '', speaker_meta)
            if re.match(r'.*\(.*?\)', speaker_meta):
                name = speaker_meta.split('(', 1)[0]
                party = speaker_meta.split('(', 1)[1].strip()[0:-1]
            else:
                name = speaker_meta
                party = ''

            line.append(name.strip())
            line.append(party.strip())

    # Delete speaker data from lines and normalize text
    for line in doc_lines:
        line[0] = normalize_string(line[0])
        if re.match(speaker_pattern, line[0]):
            line[0] = re.split(r'([.:])\s*—', line[0], 1)[2].strip()

    # Combine paragraphs
    output_data = pd.DataFrame(doc_lines,
                               columns=['text', 'coordinates', 'speechnumber', 'paragraphnumber', 'speaker',
                                        'party'])
    output_data = output_data.fillna(method='ffill')
    output_data = output_data.groupby(['speechnumber', 'paragraphnumber', 'speaker', 'party'])['text'].apply(
        ' '.join).reset_index()

    # If the spekaer is Secretário, change party speaker and add speakerrole
    output_data['speakerrole'] = output_data['speaker'].apply(lambda x: x if re.match(SPEAKER_ROLES, x) else '')
    output_data['speaker'] = output_data.apply(
        lambda x: x['party'] if re.match(SPEAKER_ROLES, x['speaker']) else x['speaker'], axis=1)
    output_data['party'] = output_data.apply(
        lambda x: '' if re.match(SPEAKER_ROLES, x['speakerrole']) else x['party'], axis=1)

    # Find index of paragraph with closing phrase and delete everything after
    closing_mask = output_data.text.str.contains(CLOSING_PHRASE, regex=True)

    try:
        closing_idx = output_data[closing_mask].index[0]
        output_data = output_data[0:closing_idx]
    except IndexError:
        print(f' Doc: {source_pdf_path.stem} Closing phrase not found')

    # Add metadata
    output_data['parliament'] = get_iso_2_digit_code(COUNTRY) + '-' + get_parliament_name(COUNTRY)
    output_data['iso3country'] = get_iso_3_digit_code(COUNTRY)
    output_data['date'] = date

    # Write parsed document
    file_name = source_pdf_path.stem
    path = year_path.joinpath(f"{file_name}_parsed.csv")
    output_data.to_csv(path, index=False)
