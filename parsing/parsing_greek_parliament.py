from docx import Document
import pathlib
import os
import re
from eia_crawling.spiders.utils import write_csv

DATE_MATCH = '.*2019.*'

SPEECH_MATCH = "^[Α-Ω]{3,}.*:"
APPLAUSE_MATCH = "^\(.*"

PRESIDENT_NAME_MATCH = "^(ΠΡΟΕΔΡΕΥΟΥΣΑ|ΠΡΟΕΔΡΕΥΩΝ).*"

USUAL_NAME_PATTERN = "^.*\(.*\)$"
def find_date_match(document):
    for paragraph in document.paragraphs:
        if re.match(DATE_MATCH, paragraph.text):
            DATE = paragraph.text
            return DATE


def find_speech_starts(document):
    # We find speech starts here
    speech_starts = []
    for i in range(len(document.paragraphs)):
        paragraph = document.paragraphs[i].text
        if re.match(SPEECH_MATCH, paragraph):
            speech_starts.append(i)

    return speech_starts


def find_speeches(document, speech_starts):
    # It's a list of lists containing paragraphs per speaker. In every list we have many paragraphs of same speaker, then we have a new list for the new speaker
    speeches = []
    for j in range(len(speech_starts)):
        # Starting location of the speech
        start_loc = speech_starts[j]

        # We take the next index as the end_loc, if it's the end of list then we take the last line of the document
        try:
            end_loc = speech_starts[j + 1]
        except IndexError:
            end_loc = len(document.paragraphs)-1

        speech = []
        for k in range(start_loc, end_loc):
            speech.append(document.paragraphs[k].text)

        speeches.append(speech)

    return speeches

def find_speaker_names(document, speech_starts):
    speaker_names = []
    for i in speech_starts:
        speaker_name = document.paragraphs[i].text.split(":")[0]
        speaker_names.append(speaker_name)

    return speaker_names

def find_period_session_sitting(document):
    PERIOD_MATCH = "^[Α-Ω]{2}΄ ΠΕΡΙΟΔΟΣ.*"
    SESSION_MATCH = "^ΣΥΝΟΔΟΣ.*"
    SITTING_MATCH = "^ΣΥΝΕΔΡΙΑΣΗ.*"

    # I initalize them all as false
    PERIOD = False
    SESSION = False
    SITTING = False

    # We will find them here
    for paragraph in document.paragraphs:

        # If we already found all, let's break
        if (PERIOD and SESSION and SITTING):
            break

        text = paragraph.text

        if re.match(PERIOD_MATCH, text):
            PERIOD = text.split(" ")[0]

        if re.match(SESSION_MATCH, text):
            SESSION = text.split(" ")[1]

        if re.match(SITTING_MATCH, text):
            SITTING = text.split(" ")[1]

    return PERIOD, SESSION, SITTING


def parse_greek_parliament(year_path: pathlib.Path,
                           year: int,
                           source_doc_path: pathlib.Path):

        document = Document(source_doc_path)
        file_name = source_doc_path.stem
        parsed_file = []

        DATE = find_date_match(document)

        # LET'S BRING IMPORTANT INFORMATION
        PERIOD, SESSION, SITTING = find_period_session_sitting(document)

        speech_starts = find_speech_starts(document)


        speeches = find_speeches(document, speech_starts)

        speaker_names = find_speaker_names(document, speech_starts)

        # Iterate over speeches
        for i in range(len(speeches)):
            # In this speech, we have many paragraphs of same speaker
            SPEECH_NUMBER = i+1
            paragraphs = speeches[i]

            SPEAKER_NAME = speaker_names[i]

            # We get the speaker role from the speaker name (if it exists) but we should do it different for the president
            try:
                # If the name has a paranthesis in the middle we skip it
                if re.match(USUAL_NAME_PATTERN, SPEAKER_NAME):
                    if re.match(PRESIDENT_NAME_MATCH, SPEAKER_NAME):
                        SPEAKER_ROLE = SPEAKER_NAME.split("(")[0][:-1]
                        SPEAKER_NAME = SPEAKER_NAME.split("(")[1][:-1]
                    else:
                        SPEAKER_ROLE = SPEAKER_NAME.split("(")[1][:-1]
                        SPEAKER_NAME = SPEAKER_NAME.split("(")[0][:-1]
                else:
                    SPEAKER_ROLE = ""
            except:
                SPEAKER_ROLE = ""
            # Let's iterate over the paragraphs
            for j in range(len(paragraphs)):

                # Get the speaker_name

                TEXT = paragraphs[j]
                PARAGRAPH_NUMBER = j+1

                # If it's the first paragraph, we should get rid of the name
                if j == 0:
                    TEXT = (":".join(TEXT.split(":")[1:]))[1:]

                if not re.match(APPLAUSE_MATCH, TEXT):

                    # Define the paragraph to be appended
                    parsed_paragraph = {"date": DATE,
                                        "agenda": "",
                                        "speechnumber": SPEECH_NUMBER,
                                        "paragraphnumber": PARAGRAPH_NUMBER,
                                        "speaker_name": SPEAKER_NAME,
                                        "party": "",
                                        "text": TEXT,
                                        "parliament": "hellenic parliament",
                                        "iso3country": "GRC",
                                        "partyname": "",
                                        "speakerrole": SPEAKER_ROLE,
                                        "period": PERIOD,
                                        "session": SESSION,
                                        "sitting": SITTING,
                                        }
                    # We append the paragraph prepared
                    parsed_file.append(parsed_paragraph)

                else:
                    continue

        # Write parsed data
        parsed_doc_path = year_path.joinpath(f"{file_name}_parsed.csv")
        write_csv(parsed_doc_path,
                  data=parsed_file,
                  fieldnames=["date",
                              "agenda",
                              "speechnumber",
                              "paragraphnumber",
                              "speaker_name",
                              "party",
                              "text",
                              "parliament",
                              "iso3country",
                              "partyname",
                              "speakerrole",
                              "period",
                              "session",
                              "sitting"])

"""
"date": "Date of the speech taken either from metadata or the file",
"agenda": "(if available) Speech topic",
"speechnumber": "Number of the speech starting from 1",
"paragraphnumber": "Paragraph number with a single speech, starting from 1, reset to 1 when new speech starts",
"speaker_name": "First and last name of the speaker",
"party": "Party name in the shortest form, usually acronym",
"text": "Text in the paragraph",
"parliament": "Parliament name in the format ISO_CODE_2-ParliamentName, e.g. DE-Bundestag",
"iso3country": "ISO 3 digit code for a country, e.g. DEU",
"partyname": "(if available) Full party name",
"speakerrole": "Role of the speaker, e.g. Ministary of Education, independent"
parliamentary_session: The name and/or number of the parliamentary session that the speech took place in. A parliamentary session includes multiple parliamentary sittings.
parliamentary_sitting: The name and/or number of the parliamentary sitting that the speech took place in.
political_party: The political party that the speaker belongs to.

"""