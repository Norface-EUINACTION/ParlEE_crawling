from docx import Document
import win32com.client as win32
import pathlib
import re
from eia_crawling.spiders.utils import get_parliament_name, get_iso_2_digit_code, get_iso_3_digit_code, write_csv, \
    normalize_string

COUNTRY = 'malta'


def save_as_docx(doc_path):
    word = win32.gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(doc_path)
    doc.Activate()
    new_path = pathlib.Path(doc_path).with_suffix('.docx')
    word.ActiveDocument.SaveAs(str(new_path), FileFormat=win32.constants.wdFormatXMLDocument)
    word.ActiveDocument.Close()


def parse_maltese_parliament(source_doc_path: pathlib,
                             year_path: pathlib.Path,
                             year: int):
    save_as_docx(str(source_doc_path))

    doc = Document(str(pathlib.Path(source_doc_path).with_suffix('.docx')))

    # Extract and remove comments to the session that appear in between speeches
    breaking_lines = []
    for p in doc.paragraphs:
        paragraph_parts = p.runs
        if all(r.italic for r in paragraph_parts) \
                and re.search(r'^("|\'|“)', p.text) is None \
                and re.search(r'(”|”\.)$', p.text) is None \
                and not all(r.text.isupper() for r in paragraph_parts):
            breaking_lines.append(p.text)
    breaking_lines = set(breaking_lines)

    paragraphs = []
    for p in doc.paragraphs:
        paragraphs.append(p.text.strip())
    paragraphs = [x for x in paragraphs if x != '' and x is not None and x not in breaking_lines]

    list_of_dicts = []
    speech_number = 1
    paragraph_number = 1
    speaker = ''
    agenda = ''

    # Delete all lines before first speech and start parsing from this idx
    start_idx = next(paragraphs.index(x) for x in paragraphs if re.search(r'^([A-Z]{2,}|\d+\.).*:.*', x) is not None)

    paragraphs = paragraphs[start_idx-1:]

    for i in range(1, len(paragraphs)):
        text = paragraphs[i]

        # Get part of the text consisting of speaker's name (and possibly role)
        if re.search(r'^([A-Z]{2,}|\d+\.).*:', paragraphs[i]) is not None:
            paragraph_number = 1
            text_split = paragraphs[i].split(':')
            if len(text_split) > 1:
                speaker = text_split[0]
                text = text_split[1]
                speech_number += 1
        else:
            paragraph_number += 1

        # Get agenda if available (for Q&A part of the session)
        if re.search(r'^\d+\.', paragraphs[i]):
            if paragraphs[i - 1].isupper():
                agenda = paragraphs[i - 1]

        # Reset agenda if parsing not Q&A speeches
        if re.search(r'^[A-Z]{2,}.*:', paragraphs[i]) and paragraphs[i - 1].isupper():
            agenda = paragraphs[i - 1]

        # Extract speaker's name and surname
        speaker = speaker.replace('.', '')
        speaker = ' '.join(x for x in speaker.split() if not x.isdigit())  # remove leading number
        if re.match(r'.*\(.*?\)', speaker):
            speaker = speaker.split('(', 1)[0]  # remove speaker role
        speaker_name = []
        for token in speaker.split():  # get only upper case leading words, as rest does not belong to name
            if token.isupper():
                speaker_name.append(token)
            else:
                break
        speaker_name = ' '.join(
            [x.lower().capitalize() for x in speaker_name if
             'onor' not in x.lower()])  # final formatting (full name must contain capitalized tokens, without title e.g. 'onor')

        list_of_dicts.append(
            {'date': pathlib.Path(source_doc_path).stem[:10],
             'parliament': get_iso_2_digit_code(COUNTRY) + '-' + get_parliament_name(COUNTRY),
             'iso3country': get_iso_3_digit_code(COUNTRY),
             'speaker': speaker_name,
             'speechnumber': speech_number-1,
             'paragraphnumber': paragraph_number,
             'agenda': agenda,
             'text': normalize_string(text)})

    # Remove texts with agenda points
    agenda_points = []
    for dct in list_of_dicts:
        agenda_points.append(dct['agenda'])
        dct['agenda'] = dct['agenda'].lower().capitalize()

    list_of_dicts = [x for x in list_of_dicts if x['text'] not in agenda_points and not x['text'].isupper()]

    # Write parsed document
    file_name = source_doc_path.stem
    path = year_path.joinpath(f"{file_name}_parsed.csv")
    write_csv(path, data=list_of_dicts, fieldnames=['date', 'parliament', 'iso3country', 'speaker', 'speechnumber',
                                                    'paragraphnumber', 'agenda', 'text'])



